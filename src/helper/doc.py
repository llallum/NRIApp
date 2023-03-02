
import docx
import re
  
# create an instance of a word document

from docx.shared import Cm    
from docx.shared import Pt
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from loguru import logger

dbgPrint = logger


name = 'Arial'

docu = docx.Document()
docu2 = docx.Document()
style2 =docu2.styles['Normal']
font = style2.font
font.name = name
font.size = Pt(10)

style = docu.styles['Normal']
font = style.font
font.name = name
font.size = Pt(11)



class WordDoc:
    def __init__(self): 
        self.doc = docx.Document()
        self.style = self.doc.styles['Normal']
        self.font = self.style.font
        self.font.name = 'Arial'
        self.font.size = Pt(11)
        self.indent = 0.4
        
    def insert_paragraph(self):
        return self.doc.add_paragraph()

    def add_run(self, text):
        self.insert_paragraph().add_run(text).bold = True
    def save(self, fullpath):
        try:
            self.doc.save(fullpath)
        except:
            dbgPrint.error("File {value} is opened.", value=fullpath)

    def title(self, title):
        self.doc.core_properties.title = title

    def author(self, author):
        self.doc.core_properties.author = author

    def traverse(self, indict):
        self.walk(indict, origin=self.doc, depth=0, indent=self.indent)

    def walk(self, indict, origin, depth=0, indent=0):
        new = indent
        if isinstance(indict, dict):
            depth = depth+1
            for k,v in indict.items():
    #            print("\t"*(indent*depth) + k)
                header = origin.add_paragraph()
                header.paragraph_format.left_indent = Cm(depth * indent)
                header.paragraph_format.space_before = Cm(0)
                header.paragraph_format.space_after = Cm(0)
                keys = re.sub('([a-zA-Z])', lambda x: x.groups()[0].upper(), k, 1)
                keys = " ".join(map(str, self.split_on_uppercase(keys, True)))
                header.add_run(keys + ": ").bold = True
                if isinstance(v, dict):
    #                header = origin.add_paragraph()
                    font = header.add_run()
                    font.font.name = name
                    header.paragraph_format.left_indent = Cm(depth * indent)
                    header.paragraph_format.space_before = Cm(0)
                    header.paragraph_format.space_after = Cm(0)
                    self.walk(v, origin, depth=depth+1, indent=indent)
                elif isinstance(v, list):
                    for i in v:
                        self.walk(i, origin,  depth+1, indent=indent)
                else:
    #                print("\t"*(indent*depth+1),  v)
                    header = origin.add_paragraph()

                    header.paragraph_format.left_indent = Cm((depth+1) * indent)
                    header.paragraph_format.space_before = Cm(0)
                    header.paragraph_format.space_after = Cm(0)
                    header.add_run(str(v))
                    #paragraphs[-1].add_run(str(val) + "  ")
        elif isinstance(indict, list):
            for i in indict:
                self.walk(i, origin,  depth+1, indent=indent)
        else:
    #        print ("\t"*(indent*depth) + indict)
            header = origin.add_paragraph()
            header.paragraph_format.left_indent = Cm(depth * indent)
            header.paragraph_format.space_before = Cm(0)
            header.paragraph_format.space_after = Cm(0)
            header.add_run(indict)


    def insertHR(self, paragraph):
        p = paragraph._p  # p is the <w:p> XML element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.insert_element_before(pBdr,
            'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
            'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
            'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
            'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
            'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
            'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
            'w:pPrChange'
        )
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)

    def split_on_uppercase(self, s, keep_contiguous=False):
        """
        Args:
            s (str): string
            keep_contiguous (bool): flag to indicate we want to 
                                    keep contiguous uppercase chars together
        Returns:

        """
        string_length = len(s)
        is_lower_around = (lambda: s[i-1].islower() or 
                           string_length > (i + 1) and s[i + 1].islower())

        start = 0
        parts = []
        for i in range(1, string_length):
            if s[i].isupper() and (not keep_contiguous or is_lower_around()):
                parts.append(s[start: i])
                start = i
        parts.append(s[start:])

        return parts